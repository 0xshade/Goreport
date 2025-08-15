#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
This is the GoReport class. GoReport handles everything from connecting to the target Gophish
server to pulling campaign information and reporting the results.
"""

try:
    # 3rd Party Libraries
    from gophish import Gophish
except:
    print("[!] Could not import the Gophish library! Make sure it is installed.\n\
Run: `python3 -m pip install gophish`\n\
Test it by running `python3` and then, in the \
Python prompt, typing `from gophish import Gophish`.")
    exit()

# Standard Libraries
import configparser
import json
import os.path
import sys
from collections import Counter
from datetime import datetime, timezone, timedelta

# 3rd Party Libraries
import requests
import xlsxwriter
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor
from requests.packages.urllib3.exceptions import InsecureRequestWarning
from user_agents import parse

requests.packages.urllib3.disable_warnings(InsecureRequestWarning)


class Goreport(object):
    """
    This class uses the Gophish library to create a new Gophish API connection
    and queries Gophish for information and results related to the specified
    campaign ID(s).
    """
    # Name of the config file -- default is ``gophish.config``
    goreport_config_file = "gophish.config"
    verbose = False

    # Variables for holding Gophish models
    results = None
    campaign = None
    timeline = None

    # Variables for holding campaign information
    cam_id = None
    cam_url = None
    cam_name = None
    cam_status = None
    launch_date = None
    created_date = None
    cam_page_name = None
    cam_smtp_host = None
    completed_date = None
    cam_redirect_url = None
    cam_from_address = None
    cam_subject_line = None
    cam_template_name = None
    cam_capturing_passwords = None
    cam_capturing_credentials = None

    # Variables and lists for tracking event numbers
    total_sent = 0
    total_opened = 0
    total_targets = 0
    total_clicked = 0
    total_reported = 0
    total_submitted = 0
    total_unique_opened = 0
    total_unique_clicked = 0
    total_unique_reported = 0
    total_unique_submitted = 0
    targets_opened = []
    targets_clicked = []
    targets_reported = []
    targets_submitted = []

    # Lists and dicts for holding prepared report data
    campaign_results_summary = []

    # Lists for holding totals for statistics
    browsers = []
    locations = []
    ip_addresses = []
    ip_and_location = {}
    operating_systems = []

    # Output options
    report_format = None
    output_word_report = None
    output_xlsx_report = None
    xlsx_header_bg_color = "#0085CA"
    xlsx_header_font_color = "#FFFFFF"

    # Timeline-specific options
    timeline_mode = False
    separate_campaigns = False

    # Timeline-specific data structures
    timeline_data = []  # List of timeline event dictionaries
    campaign_timeline_data = {}  # Dict mapping campaign_id to timeline events

    def __init__(self, report_format, config_file, google, verbose, timeline_mode=False, separate_campaigns=False, output_dir="output", rid_field=None):
        """
        Initiate the connection to the Gophish server with the provided host, port,
        and API key and prepare to use the external APIs.
        """
        try:
            # Check if an alternate config file was provided
            if config_file:
                self.goreport_config_file = config_file
            # Open the config file to make sure it exists and is readable
            config = configparser.ConfigParser()
            config.read(self.goreport_config_file)
        except Exception as e:
            print(f"[!] Could not open {self.goreport_config_file} -- make sure it exists and is readable.")
            print(f"L.. Details: {e}")
            sys.exit()

        try:
            # Read in the values from the config file
            GP_HOST = self.config_section_map(config, 'Gophish')['gp_host']
            API_KEY = self.config_section_map(config, 'Gophish')['api_key']
        except Exception as e:
            print("[!] There was a problem reading values from the gophish.config file!")
            print(f"L.. Details: {e}")
            sys.exit()

        try:
            # Read in the values from the config file
            self.IPINFO_TOKEN = self.config_section_map(config, 'ipinfo.io')['ipinfo_token']
            if not self.IPINFO_TOKEN:
                self.IPINFO_TOKEN = None
        except Exception as e:
            self.IPINFO_TOKEN = None
            print("[!] No ipinfo.io API token was found in the config. GoReport will not lookup IP addresses with ipinfo.io for additional location data.")
            print(f"L.. Details: {e}")

        try:
            # Read in the values from the config file
            self.GEOLOCATE_TOKEN = self.config_section_map(config, 'Google')['geolocate_key']
            if not self.GEOLOCATE_TOKEN:
                self.GEOLOCATE_TOKEN = None
        except Exception as e:
            self.GEOLOCATE_TOKEN = None
            if google:
                print("[!] No Google Maps API token was found in the config so GoReport will ignore the `--google` flag.")
                print(f"L.. Details: {e}")

        # Set command line options for the GoReport object
        self.google = google
        self.verbose = verbose
        self.report_format = report_format
        self.timeline_mode = timeline_mode
        self.separate_campaigns = separate_campaigns
        self.rid_field = rid_field
        # Handle output parameter - could be directory or full file path
        self.output_path = output_dir
        self.custom_filename = None

        # Check if output_dir contains a file extension (indicating it's a full path)
        if output_dir != "output" and ('.' in os.path.basename(output_dir)):
            # It's a full file path
            self.output_dir = os.path.dirname(output_dir) or "."
            self.custom_filename = os.path.basename(output_dir)
        else:
            # It's a directory
            self.output_dir = output_dir

        # Initialize timeline-specific data structures
        self.timeline_data = []
        self.campaign_timeline_data = {}

        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            print(f"[+] Created output directory: {self.output_dir}")
        # Connect to the Gophish API
        # NOTE: This step succeeds even with a bad API key, so the true test is fetching an ID
        print(f"[+] Connecting to Gophish at {GP_HOST}")
        print(f"L.. The API Authorization endpoint is: {GP_HOST}/api/campaigns/?api_key={API_KEY}")
        self.api = Gophish(API_KEY, host=GP_HOST, verify=False)

    def run(self, id_list, combine_reports, set_complete_status):
        """Run everything to process the target campaign."""
        # Output some feedback for user options will be done after we determine the final combine setting
        if set_complete_status:
            print('[+] Campaign statuses will be set to "Complete" after processing the results.')
        try:
            # Create the list of campaign IDs
            temp_id = []
            # Handle a mixed set of ranges and comma-separated IDs
            if "-" and "," in id_list:
                temp = id_list.split(",")
                for x in temp:
                    if "-" in x:
                        lower = x.split("-")[0]
                        upper = x.split("-")[1]
                        for y in range(int(lower), int(upper) + 1):
                            temp_id.append(str(y))
                    else:
                        temp_id.append(x)
            # Process IDs provided as one or more ranges
            elif "-" in id_list:
                lower = id_list.split("-")[0]
                upper = id_list.split("-")[1]
                for y in range(int(lower), int(upper) + 1):
                    temp_id.append(str(y))
            # Process single or only comma-separated IDs
            else:
                temp_id = id_list.split(",")
            id_list = temp_id
        except Exception as e:
            print("[!] Could not interpret your provided campaign IDs. \
Ensure the IDs are provided as comma-separated integers or interger ranges, e.g. 5,50-55,71.")
            print(f"L.. Details: {e}")
            sys.exit()
        # Begin processing the campaign IDs by removing any duplicates
        try:
            # Get length of user-provided list
            initial_len = len(id_list)
            # Remove duplicate IDs and sort IDs as integers
            id_list = sorted(set(id_list), key=int)
            # Get length of unique, sorted list
            unique_len = len(id_list)
        except Exception as e:
            temp = []
            for id in id_list:
                try:
                    int(id)
                except:
                    temp.append(id)
            print(f"[!] There are {len(temp)} invalid campaign ID(s), i.e. not an integer.")
            print(f"L.. Offending IDs: {','.join(temp)}")
            print(f"L.. Details: {e}")
            sys.exit()
        print(f"[+] A total of {initial_len} campaign IDs have been provided for processing.")
        # If the lengths are different, then GoReport removed one or more dupes
        if initial_len != unique_len:
            dupes = initial_len - unique_len
            print(f"L.. GoReport found {dupes} duplicate campaign IDs, so those have been trimmed.")
        # Provide  list of all IDs that will be processed
        print(f"[+] GoReport will process the following campaign IDs: {','.join(id_list)}")
        # Automatically combine reports when multiple IDs are provided
        if len(id_list) > 1 and not combine_reports:
            combine_reports = True
            print("[+] Multiple campaigns detected - automatically combining into a single report.")
        # If --combine is used with just one ID it can break reporting, so we catch that here
        elif len(id_list) == 1 and combine_reports:
            combine_reports = False
            print("[+] Single campaign - generating individual report.")
        elif len(id_list) > 1 and combine_reports:
            print("[+] Campaign results will be combined into a single report.")
        else:
            print("[+] Generating individual report for the campaign.")
        # Go through each campaign ID and get the results
        campaign_counter = 1
        for CAM_ID in id_list:
            print(f"[+] Now fetching results for Campaign ID {CAM_ID} ({campaign_counter}/{len(id_list)}).")
            try:
                # Request the details for the provided campaign ID
                self.campaign = self.api.campaigns.get(campaign_id=CAM_ID)
            except Exception as e:
                print(f"[!] There was a problem fetching this campaign {CAM_ID}'s details. Make sure your URL and API key are correct. Check HTTP vs HTTPS!")
                print(f"L.. Details: {e}")
            try:
                try:
                    # Check to see if a success message was returned with a message
                    # Possible reasons: campaign ID doesn't exist or problem with host/API key
                    if self.campaign.success is False:
                        print(f"[!] Failed to get results for campaign ID {CAM_ID}")
                        print(f"L.. Details: {self.campaign.message}")
                        # We can't let an error with an ID stop reporting, so check if this was the last ID
                        if CAM_ID == id_list[-1] and combine_reports:
                            self.finalize_timeline_data()
                            self.generate_report()
                # If self.campaign.success does not exist then we were successful
                except:
                    print("[+] Success!")
                    # Collect campaign details and process data
                    if not self.collect_all_campaign_info(combine_reports):
                        # Skip this campaign if collection failed
                        continue
                    self.process_timeline_events(combine_reports)
                    # Process timeline data for timeline reports
                    if self.timeline_mode:
                        self.process_timeline_for_report(combine_reports)
                    self.process_results(combine_reports)
                    # If the --complete flag was set, now set campaign status to Complete
                    if set_complete_status:
                        print(f"[+] Setting campaign ID {CAM_ID}'s status to Complete.")
                        try:
                            set_complete = self.api.campaigns.complete(CAM_ID)
                            try:
                                if set_complete.success is False:
                                    print(f"[!] Failed to set campaign status for ID {CAM_ID}.")
                                    print(f"L.. Details: {set_complete.message}")
                            # If set_complete.success does not exist then we were successful
                            except:
                                pass
                        except Exception as e:
                            print(f"[!] Failed to set campaign status for ID {CAM_ID}.")
                            print(f"L.. Details: {e}")
                    # Check if this is the last campaign ID in the list
                    # If this is the last ID and combined reports is on, generate the report
                    if CAM_ID == id_list[-1] and combine_reports:
                        self.finalize_timeline_data()
                        self.generate_report()
                    # Otherwise, if we are not combining reports, generate the reports
                    elif combine_reports is False:
                        self.finalize_timeline_data()
                        self.generate_report()
                    campaign_counter += 1
            except Exception as e:
                print(f"[!] There was a problem processing campaign ID {CAM_ID}!")
                print(f"L.. Details: {e}")
                sys.exit()

    def lookup_ip(self, ip):
        """Lookup the provided IP address with ipinfo.io for location data.

        Example Result:
            {'ip': '52.44.93.197',
            'hostname': 'ec2-52-44-93-197.compute-1.amazonaws.com',
            'city': 'Beaumont',
            'region': 'Texas',
            'country': 'US',
            'loc': '30.0866,-94.1274',
            'postal': '77702',
            'phone': '409',
            'org': 'AS14618 Amazon.com, Inc.'}
        """
        ipinfo_url = f"https://ipinfo.io/{ip}?token={self.IPINFO_TOKEN}"
        try:
            r = requests.get(ipinfo_url)
            return r.json()
        except Exception as e:
            print(f"[!] Failed to lookup `{ip}` with ipinfo.io.")
            print(f"L.. Details: {e}")
            return None

    def get_google_location_data(self, lat, lon):
        """Use Google's Maps API to collect location info for the provided latitude and longitude.

        Google returns a bunch of JSON with a variety of location data. This function returns
        Google's pre-formatted `formatted_address` key for a human-readable address.
        """
        google_maps_url = f"https://maps.googleapis.com/maps/api/geocode/json?latlng={lat},{lon}&sensor=false&key={self.GEOLOCATE_TOKEN}"
        r = requests.get(google_maps_url)
        maps_json = r.json()
        if r.ok:
            try:
                if "error_message" in maps_json:
                    print(f"[!] Google Maps returned an error so using Gophish coordinates. Error: {maps_json['error_message']}")
                    return f"{lat}, {lon}"
                first_result = maps_json['results'][0]
                if "formatted_address" in first_result:
                    return first_result["formatted_address"]
                # In case that key is ever unavailable try to assemble an address
                else:
                    components = first_result['address_components']
                    country = town = None
                    for c in components:
                        if "country" in c['types']:
                            country = c['long_name']
                        if "locality" in c['types']:
                            town = c['long_name']
                        if "administrative_area_level_1" in c['types']:
                            state = c['long_name']
                    return f"{town}, {state}, {country}"
            except Exception as e:
                print("[!] Failed to parse Google Maps API results so using Gophish coordinates.")
                print(f"L.. Error: {e}")
                return f"{lat}, {lon}"
        else:
            print(f"[!] Failed to contact the Google Maps API so using Gophish coordinates. Status code: {r.status_code}")
            return f"{lat}, {lon}"

    def geolocate(self, target, ipaddr, google=False):
        """Attempt to get location data for the provided target and event. Will use ipinfo.io if an
        API key is configured. Otherwise the Gophish latitude and longitude coordinates will be
        returned. If `google` is set to True this function will try to match the coordinates to a
        location using the Google Maps API.

        Returns a string: City, Region, Country
        """
        if ipaddr in self.ip_and_location:
            return self.ip_and_location[ipaddr]
        else:
            if self.IPINFO_TOKEN:
                # location_json = self.lookup_ip(event.details['browser']['address'])
                location_json = self.lookup_ip(ipaddr)
                if location_json:
                    city = region = country = "Unknown"
                    if "city" in location_json:
                        if location_json['city']:
                            city = location_json['city']
                    if "region" in location_json:
                        if location_json['region']:
                            region = location_json['region']
                    if "country" in location_json:
                        if location_json['country']:
                            country = location_json['country']
                    location = f"{city}, {region}, {country}"
                else:
                    location = f"{target.latitude}, {target.longitude}"
            elif google:
                if self.GEOLOCATE_TOKEN:
                    location = self.get_google_location_data(target.latitude, target.longitude)
                else:
                    location = f"{target.latitude}, {target.longitude}"
            else:
                location = f"{target.latitude}, {target.longitude}"
            self.locations.append(location)
            self.ip_and_location[ipaddr] = location
            return location

    def compare_ip_addresses(self, target_ip, browser_ip, verbose):
        """Compare the IP addresses of the target to that of an event. The goal: Looking for a
        mismatch that might identify some sort of interesting event. This might indicate an
        email was forwarded, a VPN was switched on/off, or maybe the target is at home.
        """
        if target_ip == browser_ip:
            return target_ip
        else:
            # We have an IP mismatch -- hard to tell why this might be.
            if verbose:
                print(f"[*] Event: This target's ({target_ip}) URL was clicked from a browser at {browser_ip}.")
            # This is an IP address not included in the results model, so we add it to our list here
            self.ip_addresses.append(browser_ip)
            return browser_ip

    def get_basic_campaign_info(self):
        """"Helper function to collect a campaign's basic details. This includes campaign name,
        status, template, and other details that are not the campaign's results.

        This keeps these calls in one place for tidiness and easier management.
        """
        self.cam_name = self.campaign.name
        self.cam_status = self.campaign.status
        self.created_date = self.campaign.created_date
        self.launch_date = self.campaign.launch_date
        self.completed_date = self.campaign.completed_date
        self.cam_url = self.campaign.url

        # Collect SMTP information
        self.smtp = self.campaign.smtp
        self.cam_from_address = self.smtp.from_address
        self.cam_smtp_host = self.smtp.host

        # Collect the template information
        self.template = self.campaign.template
        self.cam_subject_line = self.template.subject
        self.cam_template_name = self.template.name
        self.cam_template_attachments = self.template.attachments
        if self.cam_template_attachments == []:
            self.cam_template_attachments = "None Used"

        # Collect the landing page information
        self.page = self.campaign.page
        self.cam_page_name = self.page.name
        self.cam_redirect_url = self.page.redirect_url
        if self.cam_redirect_url == "":
            self.cam_redirect_url = "Not Used"
        self.cam_capturing_passwords = self.page.capture_passwords
        self.cam_capturing_credentials = self.page.capture_credentials

    def collect_all_campaign_info(self, combine_reports):
        """Collect the campaign's details and set values for each of the variables."""
        # Collect the basic campaign details
        try:
            # Check if campaign object exists
            if not self.campaign:
                print(f"[!] Campaign object is None - skipping campaign processing.")
                return False

            # Begin by checking if the ID is valid
            self.cam_id = self.campaign.id
            # Always get basic campaign info to ensure we have fresh details for each campaign
            if combine_reports and self.cam_name is None:
                print(f"[+] Reports will be combined -- setting name, dates, and URL based on campaign ID {self.cam_id}.")
            self.get_basic_campaign_info()
            # Collect the results and timeline lists
            if self.results is None:
                self.results = self.campaign.results
                self.timeline = self.campaign.timeline
            elif combine_reports:
                self.results += self.campaign.results
                self.timeline += self.campaign.timeline
            else:
                self.results = self.campaign.results
                self.timeline = self.campaign.timeline
            return True
        except Exception as e:
            print(f"[!] Error collecting campaign info: {e}")
            return False

    def process_results(self, combine_reports):
        """Process the results model to collect basic data, like total targets and event details.
        This should be run after the process_timeline_events() function which creates the
        targets_* lists.

        The results model can provide:
        first_name, last_name, email, position, and IP address
        """
        # Total length of results gives us the total number of targets
        if combine_reports and self.total_targets is None:
            self.total_targets = len(self.campaign.results)
        elif combine_reports:
            self.total_targets += len(self.campaign.results)
        else:
            # Not combining, so reset counters
            self.total_unique_opened = 0
            self.total_unique_clicked = 0
            self.total_unique_reported = 0
            self.total_unique_submitted = 0
            # Reports will not be combined, so reset tracking between reports
            self.total_targets = len(self.campaign.results)
            self.ip_addresses = []
            self.campaign_results_summary = []
        # Go through all results and extract data for statistics
        for target in self.campaign.results:
            temp_dict = {}
            # Log the IP address for additional statistics later
            if not target.ip == "":
                self.ip_addresses.append(target.ip)
                self.geolocate(target, target.ip, self.google)
            # Add all of the recipient's details and results to the temp dictionary
            temp_dict["email"] = target.email
            temp_dict["fname"] = target.first_name
            temp_dict["lname"] = target.last_name
            position = "None Provided"
            if target.position:
                position = target.position
            temp_dict["position"] = position
            temp_dict["ip_address"] = target.ip
            # Check if this target was recorded as viewing the email (tracking image)
            if target.email in self.targets_opened:
                temp_dict["opened"] = True
                self.total_unique_opened += 1
            else:
                temp_dict["opened"] = False
            # Check if this target clicked the link
            if target.email in self.targets_clicked:
                temp_dict["clicked"] = True
                self.total_unique_clicked += 1
                # Incremement the total number of opens for this target if they clicked
                # but did not display the tracking image in the email
                if target.email not in self.targets_opened:
                    self.total_unique_opened += 1
            else:
                temp_dict["clicked"] = False
            # Check if this target submitted data
            if target.email in self.targets_submitted:
                temp_dict["submitted"] = True
                self.total_unique_submitted += 1
            else:
                temp_dict["submitted"] = False
            # Check if this target reported the email
            if target.email in self.targets_reported:
                temp_dict["reported"] = True
                self.total_unique_reported += 1
            else:
                temp_dict["reported"] = False
            # Append the temp dictionary to the event summary list
            self.campaign_results_summary.append(temp_dict)

    def process_timeline_events(self, combine_reports):
        """Process the timeline model to collect basic data, like total clicks, and get detailed
        event data for recipients.

        The timeline model contains all events that occurred during the campaign.
        """
        # Create counters for enumeration
        sent_counter = 0
        click_counter = 0
        opened_counter = 0
        reported_counter = 0
        submitted_counter = 0

        # Reset target lists
        self.targets_opened = []
        self.targets_clicked = []
        self.targets_reported = []
        self.targets_submitted = []
        # Run through all events and count each of the four basic events
        for event in self.campaign.timeline:
            if event.message == "Email Sent":
                sent_counter += 1
            elif event.message == "Email Opened":
                opened_counter += 1
                self.targets_opened.append(event.email)
            elif event.message == "Clicked Link":
                click_counter += 1
                self.targets_clicked.append(event.email)
            elif event.message == "Submitted Data":
                submitted_counter += 1
                self.targets_submitted.append(event.email)
            elif event.message == "Email Reported":
                reported_counter += 1
                self.targets_reported.append(event.email)
        # Assign the counter values to our tracking lists
        if combine_reports:
            # Append, +=, totals if combining reports
            self.total_sent += sent_counter
            self.total_opened += opened_counter
            self.total_clicked += click_counter
            self.total_reported += reported_counter
            self.total_submitted += submitted_counter
        else:
            # Set tracking variables to current counter values for non-combined reports
            self.total_sent = sent_counter
            self.total_opened = opened_counter
            self.total_clicked = click_counter
            self.total_reported = reported_counter
            self.total_submitted = submitted_counter

    def extract_timeline_events(self):
        """Extract timeline events focusing on clicks and data submissions.

        Parse event details JSON to extract IP, user agent, and submitted data.
        Create structured timeline data with all required fields.
        Handle multiple clicks per user and missing data gracefully.
        """
        timeline_events = []

        # Get email sent events for context (or launch date as fallback)
        email_sent_events = {}

        for event in self.campaign.timeline:
            if event.message == "Email Sent":
                email_sent_events[event.email] = event.time

        # If no email sent events found, use launch date as fallback
        if not email_sent_events and self.launch_date:
            # Get all unique emails from click/submit events
            unique_emails = set()
            for event in self.campaign.timeline:
                if event.message in ["Clicked Link", "Submitted Data"] and event.email:
                    unique_emails.add(event.email)

            # Use launch date for all users
            for email in unique_emails:
                email_sent_events[email] = self.launch_date

        # Process click and submission events
        for event in self.campaign.timeline:
            if event.message in ["Clicked Link", "Submitted Data"]:
                try:
                    # Parse event details JSON
                    details = {}
                    if hasattr(event, 'details') and event.details:
                        try:
                            # Check if details is already a dict or needs JSON parsing
                            if isinstance(event.details, dict):
                                details = event.details
                            else:
                                details = json.loads(event.details)
                        except (json.JSONDecodeError, TypeError):
                            if self.verbose:
                                print(f"[*] Warning: Could not parse event details for {event.email}")
                            details = {}

                    # Extract browser information
                    browser_info = details.get('browser', {})
                    ip_address = browser_info.get('address', 'Unknown')
                    user_agent = browser_info.get('user-agent', 'Unknown')

                    # Extract reference ID from payload
                    payload = details.get('payload', {})
                    reference_id = 'Unknown'

                    if self.rid_field:
                        # Use custom RID field if specified
                        if self.rid_field in payload and payload[self.rid_field]:
                            reference_id = payload[self.rid_field][0] if isinstance(payload[self.rid_field], list) else payload[self.rid_field]
                    else:
                        # Check for different possible reference ID field names
                        for rid_field in ['rid', 'id', 'campaign_id', 'reference_id']:
                            if rid_field in payload and payload[rid_field]:
                                reference_id = payload[rid_field][0] if isinstance(payload[rid_field], list) else payload[rid_field]
                                break

                    # Extract submitted data for data submission events
                    submitted_data = None
                    if event.message == "Submitted Data" and payload:
                        # Remove rid from payload to get actual submitted data
                        submitted_data = {k: v for k, v in payload.items() if k != 'rid'}

                    # Create timeline event record
                    timeline_event = {
                        'campaign_id': self.cam_id,
                        'campaign_name': self.cam_name,
                        'user_email': event.email,
                        'reference_id': reference_id,
                        'email_sent_time': email_sent_events.get(event.email, 'Unknown'),
                        'event_type': event.message,
                        'event_time': event.time,
                        'submitted_data': submitted_data,
                        'ip_address': ip_address,
                        'user_agent': user_agent
                    }

                    timeline_events.append(timeline_event)

                except Exception as e:
                    if self.verbose:
                        print(f"[*] Warning: Error processing timeline event for {event.email}: {e}")
                    continue

        return timeline_events

    def format_datetime_gmt4(self, iso_datetime_str):
        """Convert ISO datetime string to human-friendly DD/MM/YYYY HH:MM format in GMT+4."""
        if not iso_datetime_str or iso_datetime_str == 'Unknown':
            return 'Unknown'

        try:
            # Handle different datetime string formats
            datetime_str = str(iso_datetime_str)

            # Replace Z with +00:00 for timezone parsing
            if datetime_str.endswith('Z'):
                datetime_str = datetime_str[:-1] + '+00:00'

            # Handle microseconds with more than 6 digits (Python limitation)
            if '.' in datetime_str and '+' in datetime_str:
                date_part, tz_part = datetime_str.split('+')
                if '.' in date_part:
                    main_part, microsec_part = date_part.split('.')
                    # Truncate microseconds to 6 digits
                    microsec_part = microsec_part[:6].ljust(6, '0')
                    datetime_str = f"{main_part}.{microsec_part}+{tz_part}"

            # Parse the ISO datetime string
            dt = datetime.fromisoformat(datetime_str)

            # Convert to GMT+4
            gmt4 = timezone(timedelta(hours=4))
            dt_gmt4 = dt.astimezone(gmt4)

            # Format as DD/MM/YYYY HH:MM
            return dt_gmt4.strftime('%d/%m/%Y %H:%M')
        except (ValueError, AttributeError) as e:
            if self.verbose:
                print(f"[*] Warning: Could not format datetime '{iso_datetime_str}': {e}")
            return str(iso_datetime_str)

    def process_timeline_for_report(self, combine_reports):
        """Process timeline data for report generation.

        Group events by user email, handle multiple clicks per user,
        correlate with email sent timestamps, and handle edge cases.
        """
        if not self.timeline_mode:
            return

        try:
            # Extract timeline events for current campaign
            timeline_events = self.extract_timeline_events()
        except Exception as e:
            if self.verbose:
                print(f"[*] Warning: Error extracting timeline events: {e}")
            timeline_events = []

        if combine_reports:
            # Store events by campaign for later processing
            # Always populate campaign_timeline_data for timeline mode
            if self.timeline_mode:
                # Use the current campaign's name directly from the campaign object
                current_campaign_name = self.campaign.name if hasattr(self.campaign, 'name') else f"Campaign {self.cam_id}"

                # Always ensure we have the structure for this campaign
                if self.cam_id not in self.campaign_timeline_data:
                    self.campaign_timeline_data[self.cam_id] = {
                        'campaign_name': '',
                        'events': [],
                        'subject': '',
                        'phish_url': '',
                        'launch_date': ''
                    }

                # Always update with the current campaign's details
                # These values are set fresh in get_basic_campaign_info() for each campaign
                self.campaign_timeline_data[self.cam_id]['campaign_name'] = current_campaign_name
                self.campaign_timeline_data[self.cam_id]['subject'] = self.cam_subject_line
                self.campaign_timeline_data[self.cam_id]['phish_url'] = self.cam_url
                self.campaign_timeline_data[self.cam_id]['launch_date'] = self.launch_date

                # Add the events
                self.campaign_timeline_data[self.cam_id]['events'].extend(timeline_events)

            # Also add to combined timeline data for non-timeline reports
            if not self.separate_campaigns:
                self.timeline_data.extend(timeline_events)
        else:
            # Single campaign - keep raw events for timeline mode
            if self.timeline_mode:
                # For timeline reports, keep raw events (one per click)
                self.timeline_data = timeline_events
            else:
                # For non-timeline reports, group by user
                self.timeline_data = timeline_events
                self._group_timeline_events_by_user(timeline_events)

    def _group_timeline_events_by_user(self, timeline_events):
        """Group timeline events by user email and aggregate multiple clicks."""
        user_timeline = {}

        for event in timeline_events:
            email = event['user_email']

            if email not in user_timeline:
                user_timeline[email] = {
                    'campaign_id': event['campaign_id'],
                    'campaign_name': event['campaign_name'],
                    'user_email': email,
                    'reference_id': event['reference_id'],
                    'email_sent_time': event['email_sent_time'],
                    'click_timestamps': [],
                    'data_submitted_time': None,
                    'submitted_data': None,
                    'ip_address': event['ip_address'],
                    'user_agent': event['user_agent']
                }

            # Add event data based on type
            if event['event_type'] == "Clicked Link":
                user_timeline[email]['click_timestamps'].append(event['event_time'])
                # Update IP and user agent from most recent click if different
                if event['ip_address'] != 'Unknown':
                    user_timeline[email]['ip_address'] = event['ip_address']
                if event['user_agent'] != 'Unknown':
                    user_timeline[email]['user_agent'] = event['user_agent']

            elif event['event_type'] == "Submitted Data":
                user_timeline[email]['data_submitted_time'] = event['event_time']
                user_timeline[email]['submitted_data'] = event['submitted_data']
                # Update IP and user agent from submission if different
                if event['ip_address'] != 'Unknown':
                    user_timeline[email]['ip_address'] = event['ip_address']
                if event['user_agent'] != 'Unknown':
                    user_timeline[email]['user_agent'] = event['user_agent']

        # Sort click timestamps for each user
        for user_data in user_timeline.values():
            user_data['click_timestamps'].sort()

        # Store grouped data (replace timeline_data with grouped version)
        if not self.separate_campaigns or not hasattr(self, '_processing_combined'):
            self.timeline_data = list(user_timeline.values())

        return user_timeline

    def finalize_timeline_data(self):
        """Finalize timeline data processing for combined campaigns."""
        if not self.timeline_mode:
            return

        if self.separate_campaigns and self.campaign_timeline_data:
            # Keep campaigns separate - data is already organized by campaign
            pass
        elif not self.separate_campaigns and self.timeline_data:
            # For timeline reports, we want to keep individual events (one row per click)
            # Just sort the events chronologically
            self.timeline_data = sorted(self.timeline_data, key=lambda x: (x.get('user_email', ''), x.get('event_time', '')))
            return

            # The code below groups events by user - skip it for timeline reports
            # Combined mode - sort all events chronologically
            # Group combined events by user across all campaigns
            combined_user_timeline = {}

            for event in self.timeline_data:
                email = event['user_email']

                if email not in combined_user_timeline:
                    combined_user_timeline[email] = {
                        'campaigns': set(),
                        'user_email': email,
                        'reference_ids': set(),
                        'email_sent_times': [],
                        'click_timestamps': [],
                        'data_submitted_times': [],
                        'submitted_data_entries': [],
                        'ip_addresses': set(),
                        'user_agents': set()
                    }

                user_data = combined_user_timeline[email]

                # Aggregate data across campaigns
                if event.get('campaign_id'):
                    user_data['campaigns'].add(f"{event['campaign_id']}: {event['campaign_name']}")
                if event.get('reference_id') and event['reference_id'] != 'Unknown':
                    user_data['reference_ids'].add(event['reference_id'])
                if event.get('email_sent_time') and event['email_sent_time'] != 'Unknown':
                    user_data['email_sent_times'].append(event['email_sent_time'])

                # Handle both raw events and grouped events
                if event.get('click_timestamps'):
                    # This is a grouped event
                    user_data['click_timestamps'].extend(event['click_timestamps'])
                elif event.get('event_type') == 'Clicked Link' and event.get('event_time'):
                    # This is a raw event
                    user_data['click_timestamps'].append(event['event_time'])

                if event.get('data_submitted_times'):
                    # This is a grouped event
                    user_data['data_submitted_times'].extend(event['data_submitted_times'])
                elif event.get('event_type') == 'Submitted Data' and event.get('event_time'):
                    # This is a raw event
                    user_data['data_submitted_times'].append(event['event_time'])

                if event.get('submitted_data'):
                    user_data['submitted_data_entries'].append(event['submitted_data'])
                if event.get('ip_address') and event['ip_address'] != 'Unknown':
                    user_data['ip_addresses'].add(event['ip_address'])
                if event.get('user_agent') and event['user_agent'] != 'Unknown':
                    user_data['user_agents'].add(event['user_agent'])

            # Convert sets to sorted lists and finalize data
            finalized_data = []
            for email, user_data in combined_user_timeline.items():
                finalized_event = {
                    'user_email': email,
                    'campaigns': ', '.join(sorted(user_data['campaigns'])),
                    'reference_ids': ', '.join(sorted(user_data['reference_ids'])) if user_data['reference_ids'] else 'Unknown',
                    'email_sent_times': sorted(user_data['email_sent_times']),
                    'click_timestamps': sorted(user_data['click_timestamps']),
                    'data_submitted_times': sorted(user_data['data_submitted_times']),
                    'submitted_data_entries': user_data['submitted_data_entries'],
                    'ip_addresses': ', '.join(sorted(user_data['ip_addresses'])) if user_data['ip_addresses'] else 'Unknown',
                    'user_agents': ', '.join(sorted(user_data['user_agents'])) if user_data['user_agents'] else 'Unknown'
                }
                finalized_data.append(finalized_event)

            # Sort by email for consistent output
            self.timeline_data = sorted(finalized_data, key=lambda x: x['user_email'])

    def generate_report(self):
        """Determines which type of report generate and the calls the appropriate reporting
        functions.
        """
        if self.report_format == "excel":
            print("[+] Building the report -- you selected a Excel/xlsx report.")
            self.output_xlsx_report = self._build_output_xlsx_file_name()
            if self.timeline_mode:
                self.write_xlsx_timeline_report()
            else:
                self.write_xlsx_report()
        elif self.report_format == "word":
            print("[+] Building the report -- you selected a Word/docx report.")
            print("[+] Looking for the template.docx to be used for the Word report.")
            if os.path.isfile("template.docx"):
                print("[+] Template was found -- proceeding with report generation...")
                print("L.. Word reports can take a while if you had a lot of recipients.")
                self.output_word_report = self._build_output_word_file_name()
                if self.timeline_mode:
                    self.write_word_timeline_report()
                else:
                    self.write_word_report()
            else:
                print("[!] Could not find the template document! Make sure 'template.docx' is in the GoReport directory.")
                sys.exit()
        elif self.report_format == "quick":
            print("[+] Quick report stats:")
            if self.timeline_mode:
                self.get_timeline_quick_stats()
            else:
                self.get_quick_stats()

    def get_quick_stats(self):
        """Present quick stats for the campaign. Just basic numbers and some details."""
        print()
        print(self.cam_name)
        print(f"Status:\t\t{self.cam_status}")
        print(f"Created:\t{self.created_date.split('T')[1].split('.')[0]} on {self.created_date.split('T')[0]}")
        print(f"Started:\t{self.launch_date.split('T')[1].split('.')[0]} on {self.launch_date.split('T')[0]}")
        if self.cam_status == "Completed":
            print(f"Completed:\t{self.completed_date.split('T')[1].split('.')[0]} on {self.completed_date.split('T')[0]}")
        print()
        print(f"Total Targets:\t{self.total_targets}")
        print(f"Emails Sent:\t{self.total_sent}")
        print(f"IPs Seen:\t{len(self.ip_addresses)}")
        print()
        print(f"Total Opened Events:\t\t{self.total_opened}")
        print(f"Total Click Events:\t\t{self.total_clicked}")
        print(f"Total Submitted Data Events:\t{self.total_submitted}")
        print()
        print(f"Individuals Who Opened:\t\t\t{self.total_unique_opened}")
        print(f"Individuals Who Clicked:\t\t{self.total_unique_clicked}")
        print(f"Individuals Who Entered Data:\t\t{self.total_unique_submitted}")
        print(f"Individuals Who Reported the Email:\t{self.total_unique_reported}")

    def get_timeline_quick_stats(self):
        """Present timeline-specific quick stats for the campaign."""
        print()
        print("=== TIMELINE REPORT ===")

        if self.separate_campaigns and self.campaign_timeline_data:
            print(f"Campaigns Processed:\t{len(self.campaign_timeline_data)}")
            print()

            for campaign_id, campaign_info in self.campaign_timeline_data.items():
                print(f"Campaign {campaign_id}: {campaign_info['campaign_name']}")
                self._print_timeline_stats_for_data(campaign_info['events'])
                print()
        else:
            campaign_name = self.cam_name if hasattr(self, 'cam_name') and self.cam_name else "Combined Campaigns"
            print(f"Timeline for: {campaign_name}")
            self._print_timeline_stats_for_data(self.timeline_data)

        # Show sample timeline events
        print("=== SAMPLE TIMELINE EVENTS ===")
        sample_events = self.timeline_data[:3] if self.timeline_data else []
        if not sample_events and self.campaign_timeline_data:
            # Get sample from first campaign
            first_campaign = next(iter(self.campaign_timeline_data.values()))
            sample_events = first_campaign['events'][:3]

        for i, event in enumerate(sample_events, 1):
            print(f"Sample {i}:")
            if isinstance(event, dict):
                print(f"  User: {event.get('user_email', 'Unknown')}")
                print(f"  Clicks: {len(event.get('click_timestamps', []))}")
                print(f"  Submitted Data: {'Yes' if event.get('submitted_data') or event.get('data_submitted_time') else 'No'}")
                print(f"  IP: {event.get('ip_address', event.get('ip_addresses', 'Unknown'))}")
            print()

    def _print_timeline_stats_for_data(self, timeline_data):
        """Print timeline statistics for a given dataset."""
        if not timeline_data:
            print("  No timeline data available")
            return

        total_users = len(timeline_data)
        users_with_clicks = 0
        users_with_submissions = 0
        total_clicks = 0
        total_submissions = 0
        unique_ips = set()

        for event in timeline_data:
            if isinstance(event, dict):
                # Count clicks
                click_timestamps = event.get('click_timestamps', [])
                if click_timestamps:
                    users_with_clicks += 1
                    if isinstance(click_timestamps, list):
                        total_clicks += len(click_timestamps)
                    else:
                        total_clicks += 1

                # Count submissions
                if (event.get('submitted_data') or
                    event.get('data_submitted_time') or
                    event.get('data_submitted_times')):
                    users_with_submissions += 1
                    submit_times = event.get('data_submitted_times', [])
                    if isinstance(submit_times, list):
                        total_submissions += len(submit_times)
                    else:
                        total_submissions += 1

                # Collect IPs
                ip_addr = event.get('ip_address', event.get('ip_addresses', ''))
                if ip_addr and ip_addr != 'Unknown':
                    if ',' in str(ip_addr):
                        unique_ips.update(str(ip_addr).split(', '))
                    else:
                        unique_ips.add(str(ip_addr))

        print(f"  Total Users:\t\t\t{total_users}")
        print(f"  Users Who Clicked:\t\t{users_with_clicks}")
        print(f"  Users Who Submitted Data:\t{users_with_submissions}")
        print(f"  Total Click Events:\t\t{total_clicks}")
        print(f"  Total Submission Events:\t{total_submissions}")
        print(f"  Unique IP Addresses:\t\t{len(unique_ips)}")

    def _build_output_xlsx_file_name(self):
        """Create the xlsx report name with output directory."""
        # Use custom filename if provided
        if hasattr(self, 'custom_filename') and self.custom_filename:
            return os.path.join(self.output_dir, self.custom_filename)

        # Generate default filename
        if self.timeline_mode and self.separate_campaigns and len(self.campaign_timeline_data) > 1:
            # For separate campaigns timeline, use generic name since each sheet has its own campaign name
            safe_name = "Multiple Campaigns"
        elif hasattr(self, 'cam_name') and self.cam_name:
            safe_name = "".join([c for c in self.cam_name if c.isalpha() or c.isdigit() or c == " "]).rstrip()
        else:
            safe_name = "Combined Campaigns"

        if self.timeline_mode:
            filename = f"Gophish Timeline Results for {safe_name}.xlsx"
        else:
            filename = f"Gophish Results for {safe_name}.xlsx"

        return os.path.join(self.output_dir, filename)

    def _build_output_word_file_name(self):
        """Create the docx report name with output directory."""
        # Use custom filename if provided
        if hasattr(self, 'custom_filename') and self.custom_filename:
            return os.path.join(self.output_dir, self.custom_filename)

        # Generate default filename
        if self.timeline_mode and self.separate_campaigns and len(self.campaign_timeline_data) > 1:
            # For separate campaigns timeline, use generic name since each sheet has its own campaign name
            safe_name = "Multiple Campaigns"
        elif hasattr(self, 'cam_name') and self.cam_name:
            safe_name = "".join([c for c in self.cam_name if c.isalpha() or c.isdigit() or c == " "]).rstrip()
        else:
            safe_name = "Combined Campaigns"

        if self.timeline_mode:
            filename = f"Gophish Timeline Results for {safe_name}.docx"
        else:
            filename = f"Gophish Results for {safe_name}.docx"

        return os.path.join(self.output_dir, filename)

    def _set_word_column_width(self, column, width):
        """Custom function for quickly and easily setting the width of a table's column in the Word
        docx output.

        This option is missing from the basic Python-docx library.
        """
        for cell in column.cells:
            cell.width = width

    def write_xlsx_report(self):
        """Assemble and output the xlsx file report.

        Throughout this function, results are assembled by adding commas and then adding to a
        results string, i.e. 'result_A' and then 'result_A' += ',result_B'. This is so the
        result can be written to the csv file and have the different pieces end up in the correct
        columns.
        """
        goreport_xlsx = xlsxwriter.Workbook(self.output_xlsx_report)
        # Bold format
        bold_format = goreport_xlsx.add_format({'bold': True})
        bold_format.set_text_wrap()
        bold_format.set_align('vcenter')
        # Centered format
        center_format = goreport_xlsx.add_format()
        center_format.set_text_wrap()
        center_format.set_align('vcenter')
        center_format.set_align('center')
        # Header format
        header_format = goreport_xlsx.add_format({'bold': True})
        header_format.set_text_wrap()
        header_format.set_align('vcenter')
        header_format.set_bg_color(self.xlsx_header_bg_color)
        header_format.set_font_color(self.xlsx_header_font_color)
        # Number cells
        num_format = goreport_xlsx.add_format()
        num_format.set_align('center')
        # Boolean cells - True
        true_format = goreport_xlsx.add_format({'bold': True})
        true_format.set_text_wrap()
        true_format.set_align('vcenter')
        true_format.set_font_color("#9C0006")
        true_format.set_bg_color("#FFC7CE")
        # Boolean cells - True
        false_format = goreport_xlsx.add_format()
        false_format.set_text_wrap()
        false_format.set_align('vcenter')
        false_format.set_font_color("#006100")
        false_format.set_bg_color("#C6EFCE")
        # Remaining cells
        wrap_format = goreport_xlsx.add_format()
        wrap_format.set_text_wrap()
        wrap_format.set_align('vcenter')

        worksheet = goreport_xlsx.add_worksheet("Overview")
        col = 0
        row = 0

        worksheet.set_column(0, 10, 62)

        worksheet.write(row, col, "Campaign Results For:", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_name}", wrap_format)
        row += 1
        worksheet.write(row, col, "Status", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_status}", wrap_format)
        row += 1
        worksheet.write(row, col, "Created", bold_format)
        worksheet.write(row, col + 1, f"{self.created_date}", wrap_format)
        row += 1
        worksheet.write(row, col, "Started", bold_format)
        worksheet.write(row, col + 1, f"{self.launch_date}", wrap_format)
        row += 1
        if self.cam_status == "Completed":
            worksheet.write(row, col, "Completed", bold_format)
            worksheet.write(row, col + 1, f"{self.completed_date}", wrap_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Campaign Details", bold_format)
        row += 1
        worksheet.write(row, col, "From", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_from_address}", wrap_format)
        row += 1
        worksheet.write(row, col, "Subject", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_subject_line}", wrap_format)
        row += 1
        worksheet.write(row, col, "Phish URL", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_url}", wrap_format)
        row += 1
        worksheet.write(row, col, "Redirect URL", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_redirect_url}", wrap_format)
        row += 1
        worksheet.write(row, col, "Attachment(s)", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_template_attachments}", wrap_format)
        row += 1
        worksheet.write(row, col, "Captured Passwords", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_capturing_credentials}", wrap_format)
        row += 1
        worksheet.write(row, col, "Stored Passwords", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_capturing_passwords}", wrap_format)
        row += 1

        worksheet.write(row, col, "")
        row += 1

        # Write a high level summary for stats
        worksheet.write(row, col, "High Level Results", bold_format)
        row += 1
        worksheet.write(row, col, "Total Targets", bold_format)
        worksheet.write(row, col + 1, self.total_targets, num_format)
        row += 1

        worksheet.write(row, col, "The following totals indicate how many events of each type Gophish recorded:", wrap_format)
        row += 1
        worksheet.write(row, col, "Total Opened Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_opened, num_format)
        row += 1
        worksheet.write(row, col, "Total Clicked Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_clicked, num_format)
        row += 1
        worksheet.write(row, col, "Total Submitted Data Events", bold_format)
        worksheet.write(row, col + 1, "", wrap_format)
        row += 1
        worksheet.write(row, col, "Total Report Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_reported, num_format)
        row += 1

        worksheet.write(row, col, "The following totals indicate how many targets participated in each event type:", wrap_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Opened", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_opened, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Clicked", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_clicked, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Submitted Data", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_submitted, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Reported", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_reported, num_format)
        row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet = goreport_xlsx.add_worksheet("Summary")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 20)

        worksheet.write(row, col, "Summary of Events", bold_format)
        row += 1

        header_col = 0
        headers = ["Email Address", "Open", "Click", "Creds", "Report", "OS", "Browser"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1

        # Sort campaign summary by each dict's email entry and then create results table
        target_counter = 0
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
        for target in ordered_results:
            worksheet.write(row, col, target['email'], wrap_format)
            if target['opened']:
                worksheet.write_boolean(row, col + 1, target['opened'], true_format)
            else:
                worksheet.write_boolean(row, col + 1, target['opened'], false_format)
            if target['clicked']:
                worksheet.write_boolean(row, col + 2, target['clicked'], true_format)
            else:
                worksheet.write_boolean(row, col + 2, target['clicked'], false_format)
            if target['submitted']:
                worksheet.write_boolean(row, col + 3, target['submitted'], true_format)
            else:
                worksheet.write_boolean(row, col + 3, target['submitted'], false_format)
            if target['reported']:
                worksheet.write_boolean(row, col + 4, target['reported'], true_format)
            else:
                worksheet.write_boolean(row, col + 4, target['reported'], false_format)
            if target['email'] in self.targets_clicked:
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 5, browser_details, wrap_format)
                        worksheet.write(row, col + 6, os_details, wrap_format)
            else:
                worksheet.write(row, col + 5, "N/A", wrap_format)
                worksheet.write(row, col + 6, "N/A", wrap_format)
            row += 1
            target_counter += 1
            print(f"[+] Created row for {target_counter} of {self.total_targets}.")

        print("[+] Finished writing events summary...")
        print("[+] Detailed results analysis is next and will take some time if you had a lot of targets...")
        # End of the event summary and beginning of the detailed results

        worksheet = goreport_xlsx.add_worksheet("Event Details")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 40)

        worksheet.write(row, col, "Detailed Analysis", bold_format)
        row += 1

        target_counter = 0
        for target in self.results:
            # Only create a Detailed Analysis section for targets with clicks
            if target.email in self.targets_clicked:
                position = ""
                if target.position:
                    position = f"({target.position})"
                worksheet.write(row, col, f"{target.first_name} {target.last_name} {position}", bold_format)
                row += 1
                worksheet.write(row, col, target.email, wrap_format)
                row += 1
                # Go through all events to find events for this target
                for event in self.timeline:
                    if event.message == "Email Sent" and event.email == target.email:
                        # Parse the timestamp into separate date and time variables
                        temp = event.time.split('T')
                        sent_date = temp[0]
                        sent_time = temp[1].split('.')[0]
                        # Record the email sent date and time in the report
                        worksheet.write(row, col, f"Sent on {sent_date.replace(',', '')} at {sent_time}", wrap_format)
                        row += 1

                    if event.message == "Email Opened" and event.email == target.email:
                        # Record the email preview date and time in the report
                        temp = event.time.split('T')
                        worksheet.write(row, col, f"Email Preview at {temp[0]} {temp[1].split('.')[0]}", wrap_format)
                        row += 1

                    if event.message == "Clicked Link" and event.email == target.email:
                        worksheet.write(row, col, "Email Link Clicked", bold_format)
                        row += 1

                        header_col = 0
                        headers = ["Time", "IP", "Location", "Browser", "Operating System"]
                        for header in headers:
                            worksheet.write(row, header_col, header, header_format)
                            header_col += 1
                        row += 1

                        temp = event.time.split('T')
                        worksheet.write(row, col, f"{temp[0]} {temp[1].split('.')[0]}", wrap_format)

                        # Check if browser IP matches the target's IP and record result
                        ip_comparison = self.compare_ip_addresses(target.ip,
                                                                  event.details['browser']['address'],
                                                                  self.verbose)
                        worksheet.write(row, col + 1, f"{ip_comparison}", wrap_format)

                        # Parse the location data
                        loc = self.geolocate(target, event.details['browser']['address'], self.google)
                        worksheet.write(row, col + 2, loc, wrap_format)

                        # Parse the user-agent string and add browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        worksheet.write(row, col + 3, browser_details, wrap_format)
                        self.browsers.append(browser_details)

                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 4, os_details, wrap_format)
                        self.operating_systems.append(os_details)
                        row += 1

                    if event.message == "Submitted Data" and event.email == target.email:
                        # Now we have events for submitted data. A few notes on this:
                        #   1. There is no expectation of a Submit event without a Clicked Link event
                        #   2. Assuming that, the following process does NOT flag IP mismatches
                        #      or add to the list of seen locations, OSs, IPs, or browsers.
                        worksheet.write(row, col, "Submitted Data Captured", bold_format)
                        row += 1

                        header_col = 0
                        headers = ["Time", "IP", "Location", "Browser", "Operating System", "Data Captured"]
                        for header in headers:
                            worksheet.write(row, header_col, header, header_format)
                            header_col += 1
                        row += 1

                        temp = event.time.split('T')
                        worksheet.write(row, col, f"{temp[0]} {temp[1].split('.')[0]}", wrap_format)

                        worksheet.write(row, col + 1, f"{event.details['browser']['address']}", wrap_format)

                        loc = self.geolocate(target, event.details['browser']['address'], self.google)
                        worksheet.write(row, col + 2, loc, wrap_format)

                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        worksheet.write(row, col + 3, browser_details, wrap_format)

                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 4, os_details, wrap_format)

                        # Get just the submitted data from the event's payload
                        submitted_data = ""
                        data_payload = event.details['payload']
                        # Get all of the submitted data
                        for key, value in data_payload.items():
                            # To get just submitted data, we drop the 'rid' key
                            if not key == "rid":
                                submitted_data += f"{key}:{str(value).strip('[').strip(']')}"
                        worksheet.write(row, col + 5, submitted_data, wrap_format)
                        row += 1

                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
            else:
                # This target had no clicked or submitted events so move on to next
                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
                continue
            worksheet.write(row, col, "")
            row += 1

        print("[+] Finished writing detailed analysis...")

        worksheet = goreport_xlsx.add_worksheet("Stats")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 35)

        worksheet.write(row, col, "Recorded Browsers Based on User-Agents:", bold_format)
        row += 1

        header_col = 0
        headers = ["Browser", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_browsers = Counter(self.browsers)
        for key, value in counted_browsers.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Record OS From Browser User-Agents:", bold_format)
        row += 1
        header_col = 0
        headers = ["Operating System", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_os = Counter(self.operating_systems)
        for key, value in counted_os.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Recorded Locations from IPs:", bold_format)
        row += 1
        header_col = 0
        headers = ["Locations", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_locations = Counter(self.locations)
        for key, value in counted_locations.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Recorded IPs:", bold_format)
        row += 1
        header_col = 0
        headers = ["IP Address", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_ip_addresses = Counter(self.ip_addresses)
        for key, value in counted_ip_addresses.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "Recorded IPs and Locations:", bold_format)
        row += 1
        header_col = 0
        headers = ["IP Address", "Location"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        for key, value in self.ip_and_location.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write(row, col + 1, f"{value}", wrap_format)
            row += 1

        goreport_xlsx.close()
        print(f"[+] Done! Check '{self.output_xlsx_report}' for your results.")

    def write_xlsx_timeline_report(self):
        """Generate Excel timeline report with timeline-specific columns."""
        timeline_xlsx = xlsxwriter.Workbook(self.output_xlsx_report)

        # Define formats
        header_format = timeline_xlsx.add_format({'bold': True})
        header_format.set_text_wrap()
        header_format.set_align('vcenter')
        header_format.set_bg_color(self.xlsx_header_bg_color)
        header_format.set_font_color(self.xlsx_header_font_color)

        wrap_format = timeline_xlsx.add_format()
        wrap_format.set_text_wrap()
        wrap_format.set_align('vcenter')

        center_format = timeline_xlsx.add_format()
        center_format.set_text_wrap()
        center_format.set_align('vcenter')
        center_format.set_align('center')

        if self.separate_campaigns and self.campaign_timeline_data:
            # Create separate sheets for each campaign
            for campaign_id, campaign_info in self.campaign_timeline_data.items():
                # Use actual campaign name for sheet name, sanitized for Excel
                campaign_name = campaign_info['campaign_name']
                clean_name = "".join([c for c in campaign_name if c.isalnum() or c in ' -_'])
                # Limit to Excel's 31 char limit
                sheet_name = clean_name[:31]
                if not clean_name:  # Fallback if name becomes empty
                    sheet_name = f"Campaign_{campaign_id}"
                worksheet = timeline_xlsx.add_worksheet(sheet_name)

                # Use raw events directly (each click on separate row)
                self._write_timeline_sheet(worksheet, campaign_info['events'], header_format, wrap_format, center_format, campaign_info, 0)
        else:
            # Combined mode - all campaigns in one sheet
            worksheet = timeline_xlsx.add_worksheet("Combined Timeline")

            if self.campaign_timeline_data:
                # Multiple campaigns combined in one sheet
                self._write_combined_timeline_sheet(worksheet, self.campaign_timeline_data, header_format, wrap_format, center_format)
            else:
                # Single campaign or fallback
                campaign_name = self.cam_name if hasattr(self, 'cam_name') and self.cam_name else "Combined Campaigns"
                # Create a campaign_info dict for consistency
                campaign_info = {
                    'campaign_name': campaign_name,
                    'subject': self.cam_subject_line if hasattr(self, 'cam_subject_line') else 'N/A',
                    'phish_url': self.cam_url if hasattr(self, 'cam_url') else 'N/A',
                    'launch_date': self.launch_date if hasattr(self, 'launch_date') else 'N/A'
                }
                self._write_timeline_sheet(worksheet, self.timeline_data, header_format, wrap_format, center_format, campaign_info, 0)

        timeline_xlsx.close()
        print(f"[+] Done! Check '{self.output_xlsx_report}' for your timeline results.")

    def _group_raw_events_by_user(self, raw_events):
        """Group raw timeline events by user email for separated campaign reports."""
        user_timeline = {}

        for event in raw_events:
            email = event['user_email']

            if email not in user_timeline:
                user_timeline[email] = {
                    'campaign_id': event['campaign_id'],
                    'campaign_name': event['campaign_name'],
                    'user_email': email,
                    'reference_id': event['reference_id'],
                    'email_sent_time': event['email_sent_time'],
                    'click_timestamps': [],
                    'data_submitted_time': None,
                    'submitted_data': None,
                    'ip_address': event['ip_address'],
                    'user_agent': event['user_agent']
                }

            # Add event data based on type
            if event['event_type'] == "Clicked Link":
                user_timeline[email]['click_timestamps'].append(event['event_time'])
                # Update IP and user agent from most recent click if different
                if event['ip_address'] != 'Unknown':
                    user_timeline[email]['ip_address'] = event['ip_address']
                if event['user_agent'] != 'Unknown':
                    user_timeline[email]['user_agent'] = event['user_agent']

            elif event['event_type'] == "Submitted Data":
                user_timeline[email]['data_submitted_time'] = event['event_time']
                user_timeline[email]['submitted_data'] = event['submitted_data']
                # Update IP and user agent from submission if different
                if event['ip_address'] != 'Unknown':
                    user_timeline[email]['ip_address'] = event['ip_address']
                if event['user_agent'] != 'Unknown':
                    user_timeline[email]['user_agent'] = event['user_agent']

        # Sort click timestamps for each user
        for user_data in user_timeline.values():
            user_data['click_timestamps'].sort()

        return list(user_timeline.values())

    def _write_combined_timeline_sheet(self, worksheet, campaign_timeline_data, header_format, wrap_format, center_format):
        """Write multiple campaigns to a single worksheet with sections for each."""
        # Set column widths (only needed once)
        worksheet.set_column(0, 0, 25)  # User Email
        worksheet.set_column(1, 1, 15)  # Reference ID
        worksheet.set_column(2, 2, 20)  # Email Sent Time
        worksheet.set_column(3, 3, 15)  # Event Type
        worksheet.set_column(4, 4, 25)  # Event Time
        worksheet.set_column(5, 5, 40)  # Submitted Data
        worksheet.set_column(6, 6, 15)  # IP Address
        worksheet.set_column(7, 7, 50)  # User Agent

        row = 0

        # Write each campaign's sections
        for campaign_id, campaign_info in campaign_timeline_data.items():
            # Debug: Print what's in campaign_info
            if self.verbose:
                print(f"[DEBUG] Writing campaign {campaign_id}:")
                print(f"  Name: {campaign_info.get('campaign_name', 'N/A')}")
                print(f"  Subject: {campaign_info.get('subject', 'N/A')}")
                print(f"  URL: {campaign_info.get('phish_url', 'N/A')}")
                print(f"  Launch: {campaign_info.get('launch_date', 'N/A')}")

            # Write this campaign's data
            row = self._write_timeline_sheet(worksheet, campaign_info['events'], header_format, wrap_format,
                                            center_format, campaign_info, row)
            # Add spacing between campaigns
            row += 3

        return row

    def _write_timeline_sheet(self, worksheet, timeline_data, header_format, wrap_format, center_format, campaign_info, start_row=0):
        """Write timeline data to a worksheet starting at specified row."""
        # Set column widths only if starting at row 0
        if start_row == 0:
            worksheet.set_column(0, 0, 25)  # User Email
            worksheet.set_column(1, 1, 15)  # Reference ID
            worksheet.set_column(2, 2, 20)  # Email Sent Time
            worksheet.set_column(3, 3, 15)  # Event Type
            worksheet.set_column(4, 4, 25)  # Event Time
            worksheet.set_column(5, 5, 40)  # Submitted Data
            worksheet.set_column(6, 6, 15)  # IP Address
            worksheet.set_column(7, 7, 50)  # User Agent

        row = start_row

        # Extract campaign details
        campaign_name = campaign_info.get('campaign_name', 'Unknown Campaign') if isinstance(campaign_info, dict) else campaign_info
        subject = campaign_info.get('subject', 'N/A') if isinstance(campaign_info, dict) else 'N/A'
        phish_url = campaign_info.get('phish_url', 'N/A') if isinstance(campaign_info, dict) else 'N/A'
        launch_date = campaign_info.get('launch_date', 'N/A') if isinstance(campaign_info, dict) else 'N/A'

        # Section 1: Campaign Details
        worksheet.write(row, 0, "CAMPAIGN DETAILS", header_format)
        worksheet.merge_range(row, 0, row, 1, "CAMPAIGN DETAILS", header_format)
        row += 1

        worksheet.write(row, 0, "Subject:", wrap_format)
        worksheet.write(row, 1, subject, wrap_format)
        row += 1

        worksheet.write(row, 0, "Phish URL:", wrap_format)
        worksheet.write(row, 1, phish_url, wrap_format)
        row += 1

        worksheet.write(row, 0, "Started:", wrap_format)
        # Format the launch date if it's available
        if launch_date != 'N/A':
            formatted_date = self.format_datetime_gmt4(launch_date) if hasattr(self, 'format_datetime_gmt4') else launch_date
            worksheet.write(row, 1, formatted_date, wrap_format)
        else:
            worksheet.write(row, 1, launch_date, wrap_format)
        row += 2

        # Section 2: Click Statistics
        worksheet.write(row, 0, "CLICK STATISTICS", header_format)
        worksheet.merge_range(row, 0, row, 1, "CLICK STATISTICS", header_format)
        row += 1

        worksheet.write(row, 0, "User", header_format)
        worksheet.write(row, 1, "Number of Clicks", header_format)
        row += 1

        # Calculate click statistics from timeline data
        click_stats = {}
        for event in timeline_data:
            if event.get('event_type') == 'Clicked Link':
                user_email = event.get('user_email', 'Unknown')
                click_stats[user_email] = click_stats.get(user_email, 0) + 1

        # Write click statistics
        if click_stats:
            for user_email, click_count in sorted(click_stats.items()):
                worksheet.write(row, 0, user_email, wrap_format)
                worksheet.write(row, 1, click_count, center_format)
                row += 1
        else:
            worksheet.write(row, 0, "No clicks recorded", wrap_format)
            row += 1

        row += 2  # Add space before timeline section

        # Section 3: Timeline Data (original section)
        # Campaign header
        worksheet.write(row, 0, f"TIMELINE DATA: {campaign_name}", header_format)
        worksheet.merge_range(row, 0, row, 7, f"TIMELINE DATA: {campaign_name}", header_format)
        row += 1

        # Column headers
        headers = [
            "User Email",
            "Reference ID",
            "Email Sent Time",
            "Event Type",
            "Event Time",
            "Submitted Data",
            "IP Address",
            "User Agent"
        ]

        # Don't add campaigns column when writing sequential campaign sections
        # The campaigns column was for the old combined mode where all events were mixed together

        for col, header in enumerate(headers):
            worksheet.write(row, col, header, header_format)
        row += 1

        # Data rows
        for event in timeline_data:
            col = 0

            # Handle both raw events and grouped events
            if isinstance(event, dict):
                # User Email
                worksheet.write(row, col, event.get('user_email', ''), wrap_format)
                col += 1

                # Reference ID
                ref_id = event.get('reference_id', event.get('reference_ids', 'Unknown'))
                worksheet.write(row, col, ref_id, center_format)
                col += 1

                # Email Sent Time
                sent_time = event.get('email_sent_time', 'Unknown')
                formatted_sent_time = self.format_datetime_gmt4(sent_time)
                worksheet.write(row, col, formatted_sent_time, center_format)
                col += 1

                # Event Type
                event_type = event.get('event_type', 'Unknown')
                worksheet.write(row, col, event_type, center_format)
                col += 1

                # Event Time
                event_time = event.get('event_time', 'Unknown')
                formatted_event_time = self.format_datetime_gmt4(event_time)
                worksheet.write(row, col, formatted_event_time, center_format)
                col += 1

                # Submitted Data
                submitted_data = event.get('submitted_data', '')
                if isinstance(submitted_data, dict):
                    submitted_data = ', '.join([f"{k}: {v}" for k, v in submitted_data.items()])
                elif not submitted_data:
                    submitted_data = 'None'
                worksheet.write(row, col, str(submitted_data), wrap_format)
                col += 1

                # IP Address
                ip_addr = event.get('ip_address', event.get('ip_addresses', 'Unknown'))
                worksheet.write(row, col, ip_addr, center_format)
                col += 1

                # User Agent
                user_agent = event.get('user_agent', event.get('user_agents', 'Unknown'))
                worksheet.write(row, col, user_agent, wrap_format)
                col += 1

                # Campaigns column removed - not needed for sequential campaign sections

                row += 1

        # Return the final row number for combined reports
        return row

    def write_word_report(self):
        """Assemble and output the Word docx file report."""
        # Create document writer using the template and a style editor
        d = Document("template.docx")
        styles = d.styles

        # Create a custom styles for table cells
        _ = styles.add_style("Cell Text", WD_STYLE_TYPE.CHARACTER)
        cell_text = d.styles["Cell Text"]
        cell_text_font = cell_text.font
        cell_text_font.name = "Calibri"
        cell_text_font.size = Pt(12)
        cell_text_font.bold = True
        cell_text_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        _ = styles.add_style("Cell Text Hit", WD_STYLE_TYPE.CHARACTER)
        cell_text_hit = d.styles["Cell Text Hit"]
        cell_text_hit_font = cell_text_hit.font
        cell_text_hit_font.name = "Calibri"
        cell_text_hit_font.size = Pt(12)
        cell_text_hit_font.bold = True
        cell_text_hit_font.color.rgb = RGBColor(0x00, 0x96, 0x00)

        _ = styles.add_style("Cell Text Miss", WD_STYLE_TYPE.CHARACTER)
        cell_text_miss = d.styles["Cell Text Miss"]
        cell_text_miss_font = cell_text_miss.font
        cell_text_miss_font.name = "Calibri"
        cell_text_miss_font.size = Pt(12)
        cell_text_miss_font.bold = True
        cell_text_miss_font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Write a campaign summary at the top of the report
        d.add_heading("Executive Summary", 1)
        p = d.add_paragraph()
        run = p.add_run(f"Campaign Results For: {self.cam_name}")
        run.bold = True
        # Runs are basically "runs" of text and must be aligned like we want
        # them aligned in the report -- thus they are pushed left
        if self.cam_status == "Completed":
            completed_status = f"Completed:\t{self.completed_date.split('T')[1].split('.')[0]} on {self.completed_date.split('T')[0]}"
        else:
            completed_status = "Still Active"
        p.add_run(f"""
Status: {self.cam_status}
Created: {self.created_date.split('T')[1].split('.')[0]} on {self.created_date.split('T')[0]}
Started: {self.launch_date.split('T')[1].split('.')[0]} on {self.launch_date.split('T')[0]}
Completed: {completed_status}

""")
        if self.cam_status == "Completed":
            print()

        # Write the campaign details -- email details and template settings
        run = p.add_run("Campaign Details")
        run.bold = True
        p.add_run(f"""
From: {self.cam_from_address}
Subject: {self.cam_subject_line}
Phish URL: {self.cam_url}
Redirect URL: {self.cam_redirect_url}
Attachment(s): {self.cam_template_attachments}
Captured Credentials: {self.cam_capturing_credentials}
Stored Passwords: {self.cam_capturing_passwords}

""")

        # Write a high level summary for stats
        run = p.add_run("High Level Results")
        run.bold = True
        p.add_run(f"""
Total Targets: {self.total_targets}

The following totals indicate how many events of each type Gophish recorded:
Total Open Events: {self.total_opened}
Total Click Events: {self.total_clicked}
Total Report Events: {self.total_reported}
Total Submitted Data Events: {self.total_submitted}

The following totals indicate how many targets participated in each event type:
Individuals Who Opened: {self.total_unique_opened}
Individuals Who Clicked: {self.total_unique_clicked}
Individuals Who Reported: {self.total_unique_reported}
Individuals Who Submitted: {self.total_unique_submitted}

""")
        d.add_page_break()

        print("[+] Finished writing high level summary...")
        # End of the campaign summary and beginning of the event summary
        d.add_heading("Summary of Events", 1)
        d.add_paragraph("The following table summarizes who opened and clicked on emails sent in this campaign.")

        # Create a table to hold the event summary results
        table = d.add_table(rows=len(self.campaign_results_summary) + 1, cols=7, style="GoReport")

        header0 = table.cell(0, 0)
        header0.text = ""
        header0.paragraphs[0].add_run("Email Address", "Cell Text").bold = True

        header1 = table.cell(0, 1)
        header1.text = ""
        header1.paragraphs[0].add_run("Open", "Cell Text").bold = True

        header2 = table.cell(0, 2)
        header2.text = ""
        header2.paragraphs[0].add_run("Click", "Cell Text").bold = True

        header3 = table.cell(0, 3)
        header3.text = ""
        header3.paragraphs[0].add_run("Data", "Cell Text").bold = True

        header4 = table.cell(0, 4)
        header4.text = ""
        header4.paragraphs[0].add_run("Report", "Cell Text").bold = True

        header5 = table.cell(0, 5)
        header5.text = ""
        header5.paragraphs[0].add_run("OS", "Cell Text").bold = True

        header6 = table.cell(0, 6)
        header6.text = ""
        header6.paragraphs[0].add_run("Browser", "Cell Text").bold = True

        # Sort campaign summary by each dict's email entry and then create results table
        target_counter = 0
        counter = 1
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
        for target in ordered_results:
            email_cell = table.cell(counter, 0)
            email_cell.text = f"{target['email']}"

            temp_cell = table.cell(counter, 1)
            if target['opened']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 2)
            if target['clicked']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 3)
            if target['submitted']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 4)
            if target['reported']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            if target['email'] in self.targets_clicked:
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        os_details = user_agent.os.family + " " + \
                            user_agent.os.version_string
                        temp_cell = table.cell(counter, 5)
                        temp_cell.text = os_details
                        temp_cell = table.cell(counter, 6)
                        temp_cell.text = browser_details
            else:
                temp_cell = table.cell(counter, 5)
                temp_cell.text = "N/A"
                temp_cell = table.cell(counter, 6)
                temp_cell.text = "N/A"
            counter += 1
            target_counter += 1
            print(f"[+] Created table entry for {target_counter} of {self.total_targets}.")

        d.add_page_break()

        # End of the event summary and beginning of the detailed results
        print("[+] Finished writing events summary...")
        print("[+] Detailed results analysis is next and may take some time if you had a lot of targets...")
        d.add_heading("Detailed Findings", 1)
        target_counter = 0
        for target in self.results:
            # Only create a Detailed Analysis section for targets with clicks
            if target.email in self.targets_clicked:
                # Create counters to track table cell locations
                opened_counter = 1
                clicked_counter = 1
                submitted_counter = 1
                # Create section starting with a header with the first and last name
                position = ""
                if target.position:
                    position = f"({target.position})"
                d.add_heading(f"{target.first_name} {target.last_name} {position}", 2)
                p = d.add_paragraph(target.email)
                p = d.add_paragraph()
                # Save a spot to record the email sent date and time in the report
                email_sent_run = p.add_run()
                # Go through all events to find events for this target
                for event in self.timeline:
                    if event.message == "Email Sent" and event.email == target.email:
                        # Parse the timestamp into separate date and time variables
                        # Ex: 2017-01-30T14:31:22.534880731-05:00
                        temp = event.time.split('T')
                        sent_date = temp[0]
                        sent_time = temp[1].split('.')[0]
                        # Record the email sent date and time in the run created earlier
                        email_sent_run.text = f"Email sent on {sent_date} at {sent_time}"

                    if event.message == "Email Opened" and event.email == target.email:
                        if opened_counter == 1:
                            # Create the Email Opened/Previewed table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Email Previews")
                            run.bold = True

                            opened_table = d.add_table(rows=1, cols=1, style="GoReport")
                            opened_table.autofit = True
                            opened_table.allow_autofit = True

                            header1 = opened_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

                        # Begin by adding a row to the table and inserting timestamp
                        opened_table.add_row()
                        timestamp = opened_table.cell(opened_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]
                        opened_counter += 1

                    if event.message == "Clicked Link" and event.email == target.email:
                        if clicked_counter == 1:
                            # Create the Clicked Link table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Email Link Clicked")
                            run.bold = True

                            clicked_table = d.add_table(rows=1, cols=5, style="GoReport")
                            clicked_table.autofit = True
                            clicked_table.allow_autofit = True

                            header1 = clicked_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

                            header2 = clicked_table.cell(0, 1)
                            header2.text = ""
                            header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

                            header3 = clicked_table.cell(0, 2)
                            header3.text = ""
                            header3.paragraphs[0].add_run("Location", "Cell Text").bold = True

                            header4 = clicked_table.cell(0, 3)
                            header4.text = ""
                            header4.paragraphs[0].add_run("Browser", "Cell Text").bold = True

                            header5 = clicked_table.cell(0, 4)
                            header5.text = ""
                            header5.paragraphs[0].add_run("Operating System",
                                                          "Cell Text").bold = True

                        clicked_table.add_row()
                        timestamp = clicked_table.cell(clicked_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        ip_add = clicked_table.cell(clicked_counter, 1)
                        # Check if browser IP matches the target's IP and record result
                        ip_add.text = self.compare_ip_addresses(
                            target.ip, event.details['browser']['address'], self.verbose)

                        # Parse the location data
                        event_location = clicked_table.cell(clicked_counter, 2)
                        event_location.text = self.geolocate(target, event.details['browser']['address'], self.google)

                        # Parse the user-agent string for browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        browser = clicked_table.cell(clicked_counter, 3)
                        browser.text = browser_details
                        self.browsers.append(browser_details)

                        op_sys = clicked_table.cell(clicked_counter, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = os_details
                        self.operating_systems.append(os_details)

                        clicked_counter += 1

                    if event.message == "Submitted Data" and event.email == target.email:
                        if submitted_counter == 1:
                            # Create the Submitted Data table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Data Captured")
                            run.bold = True

                            submitted_table = d.add_table(rows=1, cols=6, style="GoReport")
                            submitted_table.autofit = True
                            submitted_table.allow_autofit = True

                            header1 = submitted_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

                            header2 = submitted_table.cell(0, 1)
                            header2.text = ""
                            header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

                            header3 = submitted_table.cell(0, 2)
                            header3.text = ""
                            header3.paragraphs[0].add_run("Location", "Cell Text").bold = True

                            header4 = submitted_table.cell(0, 3)
                            header4.text = ""
                            header4.paragraphs[0].add_run("Browser", "Cell Text").bold = True

                            header5 = submitted_table.cell(0, 4)
                            header5.text = ""
                            header5.paragraphs[0].add_run("Operating System",
                                                          "Cell Text").bold = True

                            header6 = submitted_table.cell(0, 5)
                            header6.text = ""
                            header6.paragraphs[0].add_run("Data Captured",
                                                          "Cell Text").bold = True

                        submitted_table.add_row()
                        timestamp = submitted_table.cell(submitted_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        ip_add = submitted_table.cell(submitted_counter, 1)
                        ip_add.text = event.details['browser']['address']

                        # Parse the location data
                        event_location = submitted_table.cell(submitted_counter, 2)
                        event_location.text = self.geolocate(target, event.details['browser']['address'], self.google)

                        # Parse the user-agent string and add browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        browser = submitted_table.cell(submitted_counter, 3)
                        browser.text = browser_details

                        op_sys = submitted_table.cell(submitted_counter, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = f"{os_details}"

                        # Get just the submitted data from the event's payload
                        submitted_data = ""
                        data = submitted_table.cell(submitted_counter, 5)
                        data_payload = event.details['payload']
                        # Get all of the submitted data
                        for key, value in data_payload.items():
                            # To get just submitted data, we drop the 'rid' key
                            if not key == "rid":
                                submitted_data += f"{key}:{str(value).strip('[').strip(']')}   "
                        data.text = f"{submitted_data}"
                        submitted_counter += 1
                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")

                d.add_page_break()
            else:
                # This target had no clicked or submitted events so move on to next
                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
                continue

        print("[+] Finished writing Detailed Analysis section...")
        # End of the detailed results and the beginning of browser, location, and OS stats
        d.add_heading("Statistics", 1)
        p = d.add_paragraph("The following table shows the browsers seen:")
        # Create browser table
        browser_table = d.add_table(rows=1, cols=2, style="GoReport")
        self._set_word_column_width(browser_table.columns[0], Cm(7.24))
        self._set_word_column_width(browser_table.columns[1], Cm(3.35))

        header1 = browser_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Browser", "Cell Text").bold = True

        header2 = browser_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the operating systems seen:")

        # Create OS table
        os_table = d.add_table(rows=1, cols=2, style="GoReport")
        self._set_word_column_width(os_table.columns[0], Cm(7.24))
        self._set_word_column_width(os_table.columns[1], Cm(3.35))

        header1 = os_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Operating System", "Cell Text").bold = True

        header2 = os_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the locations seen:")

        # Create geo IP table
        location_table = d.add_table(rows=1, cols=2, style="GoReport")
        self._set_word_column_width(location_table.columns[0], Cm(7.24))
        self._set_word_column_width(location_table.columns[1], Cm(3.35))

        header1 = location_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Location", "Cell Text").bold = True

        header2 = location_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Visits", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the IP addresses captured:")

        # Create IP address table
        ip_add_table = d.add_table(rows=1, cols=2, style="GoReport")
        self._set_word_column_width(ip_add_table.columns[0], Cm(7.24))
        self._set_word_column_width(ip_add_table.columns[1], Cm(3.35))

        header1 = ip_add_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("IP Address", "Cell Text").bold = True

        header2 = ip_add_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the IP addresses matched with geolocation data:")

        # Create IP address and location table
        ip_loc_table = d.add_table(rows=1, cols=2, style="GoReport")
        self._set_word_column_width(ip_loc_table.columns[0], Cm(7.24))
        self._set_word_column_width(ip_loc_table.columns[1], Cm(3.35))

        header1 = ip_loc_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("IP Address", "Cell Text").bold = True

        header2 = ip_loc_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Location", "Cell Text").bold = True

        # Counters are used here again to track rows
        counter = 1
        # Counter is used to count all elements in the lists to create a unique list with totals
        counted_browsers = Counter(self.browsers)
        for key, value in counted_browsers.items():
            browser_table.add_row()
            cell = browser_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = browser_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_os = Counter(self.operating_systems)
        for key, value in counted_os.items():
            os_table.add_row()
            cell = os_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = os_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_locations = Counter(self.locations)
        for key, value in counted_locations.items():
            location_table.add_row()
            cell = location_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = location_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_ip_addresses = Counter(self.ip_addresses)
        for key, value in counted_ip_addresses.items():
            ip_add_table.add_row()
            cell = ip_add_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = ip_add_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        for key, value in self.ip_and_location.items():
            ip_loc_table.add_row()
            cell = ip_loc_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = ip_loc_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        # Finalize document and save it as the value of output_word_report
        d.save(f"{self.output_word_report}")
        print(f"[+] Done! Check \"{self.output_word_report}\" for your results.")

    def write_word_timeline_report(self):
        """Generate Word timeline report with timeline-specific content."""
        # Create document writer using the template
        d = Document("template.docx")

        # Add timeline report header
        d.add_heading("Timeline Report", 1)

        if self.separate_campaigns and self.campaign_timeline_data:
            # Create sections for each campaign
            for campaign_id, campaign_info in self.campaign_timeline_data.items():
                d.add_heading(f"Campaign {campaign_id}: {campaign_info['campaign_name']}", 2)
                self._write_timeline_word_section(d, campaign_info['events'])
                d.add_page_break()
        else:
            # Single section for combined data
            campaign_name = self.cam_name if hasattr(self, 'cam_name') and self.cam_name else "Combined Campaigns"
            d.add_heading(f"Timeline for {campaign_name}", 2)
            self._write_timeline_word_section(d, self.timeline_data)

        # Save the document
        d.save(f"{self.output_word_report}")
        print(f"[+] Done! Check \"{self.output_word_report}\" for your timeline results.")

    def _write_timeline_word_section(self, document, timeline_data):
        """Write timeline data section to Word document."""
        if not timeline_data:
            document.add_paragraph("No timeline data available.")
            return

        # Add summary paragraph
        document.add_paragraph(f"This section contains timeline data for {len(timeline_data)} users.")

        # Create timeline table
        table = document.add_table(rows=1, cols=8)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        headers = [
            "User Email",
            "Reference ID",
            "Email Sent Time",
            "Click Timestamps",
            "Data Submitted Time",
            "Submitted Data",
            "IP Address",
            "User Agent"
        ]

        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for event in timeline_data:
            row_cells = table.add_row().cells

            # Handle both raw events and grouped events
            if isinstance(event, dict):
                # User Email
                row_cells[0].text = event.get('user_email', '')

                # Reference ID
                ref_id = event.get('reference_id', event.get('reference_ids', 'Unknown'))
                row_cells[1].text = str(ref_id)

                # Email Sent Time
                sent_time = event.get('email_sent_time', '')
                if isinstance(event.get('email_sent_times'), list) and event['email_sent_times']:
                    sent_time = ', '.join(event['email_sent_times'])
                row_cells[2].text = str(sent_time)

                # Click Timestamps
                click_times = event.get('click_timestamps', [])
                if isinstance(click_times, list):
                    click_times_str = ', '.join(click_times) if click_times else 'None'
                else:
                    click_times_str = str(click_times) if click_times else 'None'
                row_cells[3].text = click_times_str

                # Data Submitted Time
                submit_time = event.get('data_submitted_time', '')
                if isinstance(event.get('data_submitted_times'), list) and event['data_submitted_times']:
                    submit_time = ', '.join(event['data_submitted_times'])
                row_cells[4].text = str(submit_time) if submit_time else 'None'

                # Submitted Data
                submitted_data = event.get('submitted_data', event.get('submitted_data_entries', ''))
                if isinstance(submitted_data, list) and submitted_data:
                    # Convert list of dicts to readable format
                    data_strs = []
                    for data_entry in submitted_data:
                        if isinstance(data_entry, dict):
                            data_strs.append(', '.join([f"{k}: {v}" for k, v in data_entry.items()]))
                        else:
                            data_strs.append(str(data_entry))
                    submitted_data = '; '.join(data_strs)
                elif isinstance(submitted_data, dict):
                    submitted_data = ', '.join([f"{k}: {v}" for k, v in submitted_data.items()])
                elif not submitted_data:
                    submitted_data = 'None'
                row_cells[5].text = str(submitted_data)

                # IP Address
                ip_addr = event.get('ip_address', event.get('ip_addresses', 'Unknown'))
                row_cells[6].text = str(ip_addr)

                # User Agent
                user_agent = event.get('user_agent', event.get('user_agents', 'Unknown'))
                row_cells[7].text = str(user_agent)

        # Add some spacing after the table
        document.add_paragraph("")

    def config_section_map(self, config_parser, section):
        """This function helps by reading accepting a config file section, from gophish.config,
        and returning a dictionary object that can be referenced for configuration settings.
        """
        section_dict = {}
        options = config_parser.options(section)
        for option in options:
            try:
                section_dict[option] = config_parser.get(section, option)
                if section_dict[option] == -1:
                    print(f"[-] Skipping: {option}")
            except:
                print(f"[!] There was an error with: {option}")
                section_dict[option] = None
        return section_dict
